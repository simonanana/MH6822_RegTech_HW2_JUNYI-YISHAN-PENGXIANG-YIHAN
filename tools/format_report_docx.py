from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips


USABLE_WIDTH_DXA = 9746  # US Letter width minus 2.2 cm left/right margins.
HEADER_FILL = "D9EAF7"
ALT_ROW_FILL = "F7FBFF"
CODE_FILL = "F6F8FA"
GRID_LINE = "B7C9D6"
CODE_LINE = "D0D7DE"
INK = RGBColor(31, 41, 55)
BLUE = RGBColor(15, 76, 110)
GREY = RGBColor(107, 114, 128)
GREY_DARK = RGBColor(55, 65, 81)

SUMMARY_TABLE_HEADERS = ("data quality status", "count")


def paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def set_font(run, *, name: str, size: float, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_style_font(style, *, name: str, size: float, color: RGBColor, bold: bool = False) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def set_style_spacing(style, *, before: float, after: float, line_spacing: float, keep_next: bool = False) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    fmt.keep_with_next = keep_next
    fmt.keep_together = keep_next


def first_child(parent, tag: str):
    return parent.find(qn(tag))


def remove_children(parent, tag: str) -> None:
    for child in list(parent):
        if child.tag == qn(tag):
            parent.remove(child)


def append_page_field(paragraph, field_name: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_name} "
    paragraph.add_run()._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)

    value = paragraph.add_run("1")
    set_font(value, name="Times New Roman", size=8, color=GREY)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    remove_children(tc_pr, "w:shd")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    cell.width = Twips(width_dxa)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = first_child(tc_pr, "w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.insert(0, tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    remove_children(tbl_pr, "w:tblBorders")
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), GRID_LINE)
        borders.append(border)
    tbl_pr.append(borders)


def set_table_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    remove_children(tbl_pr, "w:tblCellMar")
    margins = OxmlElement("w:tblCellMar")
    for name, width in {"top": "62", "bottom": "62", "left": "110", "right": "110"}.items():
        margin = OxmlElement(f"w:{name}")
        margin.set(qn("w:w"), width)
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    tbl_pr.append(margins)


def set_table_width(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = first_child(tbl_pr, "w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(USABLE_WIDTH_DXA))
    remove_children(tbl_pr, "w:tblLayout")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    table.autofit = False


def set_table_grid(table, widths: list[int]) -> None:
    for child in list(table._tbl):
        if child.tag == qn("w:tblGrid"):
            table._tbl.remove(child)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    insert_at = 1 if len(table._tbl) and table._tbl[0].tag == qn("w:tblPr") else 0
    table._tbl.insert(insert_at, grid)


def keep_row_together(row, *, repeat_header: bool = False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    remove_children(tr_pr, "w:trHeight")
    if first_child(tr_pr, "w:cantSplit") is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    if repeat_header and first_child(tr_pr, "w:tblHeader") is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def keep_short_table_together(table) -> None:
    if len(table.rows) > 7:
        return
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = False
                paragraph.paragraph_format.keep_with_next = row_index < len(table.rows) - 1


def is_numeric_like(value: str) -> bool:
    stripped = value.strip().replace(",", "")
    return bool(stripped) and all(ch.isdigit() or ch == "." for ch in stripped)


def table_column_widths(headers: list[str]) -> list[int]:
    key = tuple(header.lower() for header in headers)
    if key == ("student id", "name", "responsibility", "main deliverables", "related module"):
        return [1120, 1320, 2260, 3970, 1076]
    if key == ("trade", "data quality / scope", "main findings", "interpretation"):
        return [780, 1700, 3550, 3716]
    if key == ("trade", "platform", "function", "reporting issue", "engine conclusion"):
        return [900, 1500, 1900, 3350, 2096]
    if key == ("trade", "q1 exposure", "q2 hedge substitute", "q3 price signal", "conclusion"):
        return [900, 1800, 2400, 2400, 2246]
    if key == ("trade", "no upi", "non-iso ccy", "id gaps", "ec-1 threshold", "action flag"):
        return [800, 950, 1200, 900, 1500, 4396]
    if key == ("trade id", "scenario", "expected result"):
        return [1000, 4500, 4246]
    if key == ("attribute", "purpose"):
        return [2600, 7146]
    if len(headers) == 2:
        return [6200, 3546]
    if len(headers) == 3:
        return [3000, 3750, 2996]
    base = USABLE_WIDTH_DXA // max(len(headers), 1)
    widths = [base] * len(headers)
    widths[-1] += USABLE_WIDTH_DXA - sum(widths)
    return widths


def cell_text(cell) -> str:
    return " ".join(paragraph.text.strip() for paragraph in cell.paragraphs).strip()


def format_tables(doc: Document) -> None:
    center_headers = {
        "count",
        "trades",
        "q1 exposure",
        "q2 hedge substitute",
        "q3 price signal",
        "no upi",
        "non-iso ccy",
        "id gaps",
        "ec-1 threshold",
    }

    for table in doc.tables:
        if not table.rows:
            continue
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_table_width(table)
        set_table_borders(table)
        set_table_cell_margins(table)
        headers = [cell_text(cell) for cell in table.rows[0].cells]
        widths = table_column_widths(headers)
        set_table_grid(table, widths)
        header_key = tuple(header.lower() for header in headers)
        is_contribution_table = tuple(header.lower() for header in headers) == (
            "student id",
            "name",
            "responsibility",
            "main deliverables",
            "related module",
        )
        is_summary_table = header_key == SUMMARY_TABLE_HEADERS
        table_font_size = 8.5 if is_contribution_table else 9
        keep_short_table_together(table)

        for row_index, row in enumerate(table.rows):
            keep_row_together(row, repeat_header=row_index == 0)
            if is_summary_table:
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                row.height = Pt(31 if row_index in {0, 4, 8} else 27)
            for col_index, cell in enumerate(row.cells):
                width = widths[min(col_index, len(widths) - 1)]
                set_cell_width(cell, width)
                if row_index == 0:
                    set_cell_shading(cell, HEADER_FILL)
                elif row_index % 2 == 0:
                    set_cell_shading(cell, ALT_ROW_FILL)

                header = headers[col_index].lower() if col_index < len(headers) else ""
                if is_contribution_table and (row_index == 0 or header in {"student id", "name", "related module"}):
                    align = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    align = WD_ALIGN_PARAGRAPH.CENTER if header in center_headers or is_numeric_like(cell_text(cell)) else WD_ALIGN_PARAGRAPH.LEFT
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = (
                        1.08 if is_summary_table else (1.0 if is_contribution_table else 1.05)
                    )
                    paragraph.alignment = align
                    for run in paragraph.runs:
                        set_font(
                            run,
                            name="Times New Roman",
                            size=12.3 if is_summary_table and row_index in {0, 4, 8} else (12 if is_summary_table else table_font_size),
                            color=BLUE if row_index == 0 else INK,
                            bold=row_index == 0 or (is_summary_table and row_index in {4, 8}),
                        )


def add_paragraph_border(paragraph, *, top: bool, bottom: bool) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    remove_children(p_pr, "w:pBdr")
    borders = OxmlElement("w:pBdr")
    for name, enabled in {"top": top, "left": True, "bottom": bottom, "right": True}.items():
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "single" if enabled else "nil")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), CODE_LINE)
        borders.append(border)
    p_pr.append(borders)


def add_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    remove_children(p_pr, "w:shd")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def set_paragraph_element_spacing(p_elm, *, before: int | None = None, after: int | None = None) -> None:
    p_pr = first_child(p_elm, "w:pPr")
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        p_elm.insert(0, p_pr)
    spacing = first_child(p_pr, "w:spacing")
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    if before is not None:
        spacing.set(qn("w:before"), str(before))
    if after is not None:
        spacing.set(qn("w:after"), str(after))


def format_table_spacing(doc: Document) -> None:
    body = doc._body._element
    children = list(body)
    for index, child in enumerate(children):
        if child.tag != qn("w:tbl"):
            continue
        for prev in reversed(children[:index]):
            if prev.tag == qn("w:p"):
                set_paragraph_element_spacing(prev, after=80)
                break
        for nxt in children[index + 1 :]:
            if nxt.tag == qn("w:p"):
                set_paragraph_element_spacing(nxt, before=110)
                break


def is_code_paragraph(paragraph) -> bool:
    style = paragraph.style
    style_id = (style.style_id or "").lower() if style is not None else ""
    style_name = (style.name or "").lower() if style is not None else ""
    return style_id in {"sourcecode", "codeblock"} or style_name in {"source code", "sourcecode", "code block"}


def format_code_blocks(doc: Document) -> None:
    code_block: list = []
    for paragraph in doc.paragraphs:
        if is_code_paragraph(paragraph):
            code_block.append(paragraph)
            continue
        if code_block:
            apply_code_block(code_block)
            code_block = []
    if code_block:
        apply_code_block(code_block)


def format_inline_code_runs(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    for paragraph in paragraphs:
        if is_code_paragraph(paragraph):
            continue
        for run in paragraph.runs:
            style = run.style
            style_id = (style.style_id or "").lower() if style is not None else ""
            style_name = (style.name or "").lower() if style is not None else ""
            if style_id == "verbatimchar" or style_name == "verbatim char":
                set_font(run, name="Consolas", size=9, color=INK, bold=False)


def is_list_paragraph(paragraph) -> bool:
    p_pr = paragraph._p.pPr
    return p_pr is not None and p_pr.numPr is not None


def format_paragraph_density(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        style_name = (paragraph.style.name or "") if paragraph.style is not None else ""
        text = paragraph.text.strip()
        if is_code_paragraph(paragraph):
            continue
        if style_name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            continue
        if text.startswith("Appendix "):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            if text.startswith("Appendix B:"):
                paragraph.paragraph_format.space_before = Pt(9)
            continue
        if is_list_paragraph(paragraph):
            paragraph.paragraph_format.left_indent = Cm(0.55)
            paragraph.paragraph_format.first_line_indent = Cm(-0.25)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2.5)
            paragraph.paragraph_format.line_spacing = 1.08
            paragraph.paragraph_format.keep_together = False
            paragraph.paragraph_format.keep_with_next = False
            continue
        paragraph.paragraph_format.line_spacing = 1.1
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.keep_together = False
        paragraph.paragraph_format.keep_with_next = False


def format_callouts(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Key takeaway:"):
            add_paragraph_shading(paragraph, "EAF4FB")
            add_paragraph_border(paragraph, top=True, bottom=True)
            paragraph.paragraph_format.left_indent = Cm(0.16)
            paragraph.paragraph_format.right_indent = Cm(0.16)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(7)
            paragraph.paragraph_format.line_spacing = 1.08
            for run in paragraph.runs:
                if run.text.startswith("Key takeaway"):
                    set_font(run, name="Times New Roman", size=10.5, color=BLUE, bold=True)
                else:
                    set_font(run, name="Times New Roman", size=10.5, color=INK)


def apply_code_block(paragraphs: list) -> None:
    for index, paragraph in enumerate(paragraphs):
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.left_indent = Cm(0.18)
        paragraph.paragraph_format.right_indent = Cm(0.18)
        paragraph.paragraph_format.keep_together = False
        paragraph.paragraph_format.keep_with_next = False
        add_paragraph_shading(paragraph, CODE_FILL)
        add_paragraph_border(paragraph, top=index == 0, bottom=index == len(paragraphs) - 1)
        for run in paragraph.runs:
            set_font(run, name="Consolas", size=8.5, color=INK, bold=False)
            run.font.no_proof = True


def format_styles(doc: Document) -> None:
    style_specs = {
        "Normal": ("Times New Roman", 10.5, INK, False, 0, 5, 1.1, False),
        "Body Text": ("Times New Roman", 10.5, INK, False, 0, 5, 1.1, False),
        "First Paragraph": ("Times New Roman", 10.5, INK, False, 0, 5, 1.1, False),
        "Heading 1": ("Times New Roman", 19, BLUE, True, 0, 8, 1.08, True),
        "Heading 2": ("Times New Roman", 13, BLUE, True, 12, 6, 1.08, True),
        "Heading 3": ("Times New Roman", 12, GREY_DARK, True, 9, 4, 1.08, True),
        "Source Code": ("Consolas", 8.5, INK, False, 0, 0, 1.0, False),
    }
    for style_name, (font, size, color, bold, before, after, line_spacing, keep_next) in style_specs.items():
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        set_style_font(style, name=font, size=size, color=color, bold=bold)
        set_style_spacing(style, before=before, after=after, line_spacing=line_spacing, keep_next=keep_next)


def format_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)
        section.top_margin = Cm(2.24)
        section.bottom_margin = Cm(1.98)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.header_distance = Cm(0.9)
        section.footer_distance = Cm(0.9)

        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0]
        paragraph.clear()
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(17.1), WD_TAB_ALIGNMENT.RIGHT)
        left = paragraph.add_run("MH6822 ASS2 | OTC Trade Reporting Compliance Engine")
        set_font(left, name="Times New Roman", size=8, color=GREY)
        paragraph.add_run("\t")
        right = paragraph.add_run("RegTech Series")
        set_font(right, name="Times New Roman", size=8, color=GREY)

        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("Page ")
        set_font(run, name="Times New Roman", size=8, color=GREY)
        append_page_field(paragraph, "PAGE")
        run = paragraph.add_run(" of ")
        set_font(run, name="Times New Roman", size=8, color=GREY)
        append_page_field(paragraph, "NUMPAGES")


def improve_pagination(doc: Document) -> None:
    page_break_before = {
        "Appendix A: Team Contribution Statement",
        "Appendix C: Proposed EventContract JSON Schema",
    }
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text in page_break_before:
            before = paragraph.insert_paragraph_before()
            before.add_run().add_break(WD_BREAK.PAGE)
        style_name = (paragraph.style.name or "") if paragraph.style is not None else ""
        if text.startswith("Appendix ") or style_name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True


def format_docx(input_path: Path, output_path: Path) -> None:
    doc = Document(input_path)
    format_styles(doc)
    improve_pagination(doc)
    format_tables(doc)
    format_inline_code_runs(doc)
    format_paragraph_density(doc)
    format_table_spacing(doc)
    format_callouts(doc)
    format_code_blocks(doc)
    format_sections(doc)

    if not doc.sections:
        doc.add_section(WD_SECTION.NEW_PAGE)
    doc.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Apply final Word formatting to the ASS2 report DOCX.")
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("--out", type=Path, default=None, help="Output DOCX path. Defaults to in-place update.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output = args.out or args.input_docx
    original_mode = args.input_docx.stat().st_mode & 0o777
    format_docx(args.input_docx, output)
    output.chmod(original_mode)
    print(output)


if __name__ == "__main__":
    main()
